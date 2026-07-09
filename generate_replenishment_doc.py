#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成RDC订单满足分析看板 - 补货建议模块设计文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading_elm.set(qn('w:shd'), qn('w:fill'), color)

def add_styled_heading(doc, text, level=1):
    """添加格式化标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return heading

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def main():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ==== Title Page ====
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('RDC订单满足分析看板\n补货建议模块 — 设计文档')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('文档版本：v1.0\n编制日期：2026年7月2日\n编制人：张宇飞 — 补货计划经理')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_page_break()

    # ==== TOC ====
    add_styled_heading(doc, '目录', level=1)
    toc_items = [
        '1. 模块概述',
        '2. 每日补货建议',
        '   2.1 算法流程',
        '   2.2 综合评分公式',
        '   2.3 建议补货量计算',
        '   2.4 缺货原因分类与应对',
        '   2.5 紧急度分级',
        '3. 中期补货策略调整',
        '   3.1 分仓需求调整',
        '   3.2 安全库存调整',
        '   3.3 异常模式检测',
        '4. 数据依赖关系',
        '5. 页面交互设计',
        '6. 附录：算法公式汇总',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ==== Chapter 1: 模块概述 ====
    add_styled_heading(doc, '1. 模块概述', level=1)
    
    doc.add_paragraph(
        '补货建议模块是RDC订单满足分析看板的核心决策支持功能，旨在基于现有缺货数据、'
        '历史订单数据、缺货原因维护记录等多维度信息，自动生成补货建议，帮助补货计划经理'
        '快速识别需要重点关注的SKU和RDC组合，并提供量化补货量建议。'
    )
    
    doc.add_paragraph('模块包含两个子模块：', style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('每日补货建议（单日快照）')
    run.bold = True
    p.add_run('：假设当天需要补货，基于当日缺货数据、近30天历史订单趋势、'
              '缺货原因记录等，给出当天需要重点关注的Top 20 SKU及建议补货量。')

    p = doc.add_paragraph()
    run = p.add_run('中期补货策略调整（滚动更新）')
    run.bold = True
    p.add_run('：基于30天历史数据的统计分析，识别分仓需求偏差、安全库存不足、'
              '异常需求趋势等中长期问题，为补货参数调整提供数据支持。')

    # ==== Chapter 2: 每日补货建议 ====
    add_styled_heading(doc, '2. 每日补货建议', level=1)
    
    add_styled_heading(doc, '2.1 算法流程', level=2)
    doc.add_paragraph(
        '每日补货建议基于以下步骤自动生成，固定选取最新日期（当前数据中最新的一天）'
        '作为分析基准日：'
    )
    
    steps = [
        ('Step 1 — 数据筛选', 
         '从缺货汇总（dataStore.shortage）中筛选当日的所有SKU×RDC缺货组合。'
         '排除条件：总仓缺货 且 在途量=0（标记为"需总仓协调"，不纳入补货建议范围）。'),
        ('Step 2 — 计算日均需求',
         '从订单明细（dataStore.orderDetail）中提取近30天的订单数据，'
         '按 SKU×RDC 组合汇总，求日均订单支数。'),
        ('Step 3 — 计算综合评分',
         '对每个候选SKU计算综合评分（0-100分），评分公式见2.2节。'
         '按评分降序排列，取Top 20作为重点关注列表。'),
        ('Step 4 — 计算建议补货量',
         '基于日均需求、补货周期（4天）、大仓库存、在途量计算建议补货量，公式见2.3节。'),
        ('Step 5 — 匹配缺货原因',
         '从缺货原因维护记录（_shortageReasons）中读取该SKU×RDC的原因标签，'
         '匹配对应应对措施建议。'),
    ]
    
    for title, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title + '：')
        run.bold = True
        p.add_run(desc)

    add_styled_heading(doc, '2.2 综合评分公式', level=2)
    
    doc.add_paragraph('综合评分旨在将多个维度的信息融合为一个0-100分的可排序指标。'
                      '各维度权重基于快消品补货管理行业经验设定：')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Score = (1-R)*0.25 + Trend*0.15 + ABC*0.15 + RDC*0.10 + Short*0.35'
    )
    run.font.size = Pt(12)
    run.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        '其中：\n'
        '• R = 近30天该SKU在该RDC的订单满足率 = 总履约支数 / 订单总支数\n'
        '  贡献 (1-R)*0.25：满足率越低，得分越高\n\n'
        '• Trend = 近7天缺货趋势归一化值\n'
        '  若近7天日均缺货量 > 前7天日均缺货量，则 Trend = 1（上升趋势）\n'
        '  否则 Trend = 0\n'
        '  贡献 Trend*0.15：缺货在加重则得分更高\n\n'
        '• ABC = ABC分类系数\n'
        '  A类 = 1.0, B类 = 0.6, C类 = 0.3\n'
        '  贡献 ABC*0.15：A类SKU（高价值/高销量）得分更优先\n\n'
        '• RDC = 涉及RDC数量归一化值 = affectedRdcCount / 6\n'
        '  贡献 RDC*0.10：影响面越广越需要关注\n\n'
        '• Short = 缺货量归一化值 = min(shortBoxes / maxShortBoxes, 1)\n'
        '  maxShortBoxes = 当前分析批次中最大的缺货箱数\n'
        '  贡献 Short*0.35：缺货量是最重要的直接指标'
    )

    add_styled_heading(doc, '2.3 建议补货量计算', level=2)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        '建议补货量 = max(0, 日均需求 × 4天 − 大仓库存箱数 − 在途箱数)'
    )
    run.font.size = Pt(12)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(
        '参数说明：\n'
        '• 日均需求 = 近30天该SKU在对应RDC上的日均订单支数\n'
        '• 补货周期 = 4天（RDC补货的标准周期，按宇飞实际经验设为3-4天，取4天保守估计）\n'
        '• 大仓库存箱数 = 缺货汇总中的 dcStock 字段（大仓库存箱数）\n'
        '• 在途箱数 = 缺货汇总中的 rdcTransitTotal 字段（RDC在途总箱数）\n\n'
        '若计算结果为负数，表示当前库存+在途已能满足需求，建议量为0。'
    )

    add_styled_heading(doc, '2.4 缺货原因分类与应对建议', level=2)
    
    add_table(doc,
        ['缺货原因', '含义', '应对建议', '紧急度'],
        [
            ['总仓缺货', '大仓库存箱数 <= 1', '标记为"需总仓协调"，不纳入补货建议；同步提示大仓补货', '🔴 紧急'],
            ['在途缺货', '大仓有库存但RDC缺货', '建议紧急调拨，补货周期3-4天', '🔴 紧急'],
            ['库存在KA库', '库存在KA专有库中未释放', '协调KA库释放库存，或从共享库补货', '🟡 重要'],
            ['低于2/3效期', '库存效期不足2/3不可发货', '建议调拨替代批次或协调客户接受', '🟡 重要'],
            ['指定效期发货', '客户要求特定效期', '建议客户沟通，协商效期放宽', '🔵 关注'],
            ['入库时间差', '库存在途即将入库', '短期可自行解决，持续监控', '🔵 关注'],
            ['缺货原因为空', '尚未维护缺货原因', '标记为"待确认"，提示尽快维护缺货原因', '🟡 重要'],
        ]
    )

    add_styled_heading(doc, '2.5 紧急度分级', level=2)
    
    add_table(doc,
        ['级别', '评分范围', '标识色', '含义'],
        [
            ['紧急', 'Score >= 70', '红色 (EF4444)', '需立即处理，影响多个RDC或A类SKU，缺货量大'],
            ['重要', '40 <= Score < 70', '黄色 (F59E0B)', '需要关注，近几天可能进一步恶化'],
            ['关注', 'Score < 40', '蓝色 (2563EB)', '持续监控，按正常补货周期处理即可'],
        ]
    )

    doc.add_page_break()

    # ==== Chapter 3: 中期补货策略 ====
    add_styled_heading(doc, '3. 中期补货策略调整', level=1)
    
    add_styled_heading(doc, '3.1 分仓需求调整', level=2)
    
    doc.add_paragraph(
        '目标：识别各RDC之间的需求比例偏差，建议调整分仓分配参数，使补货流向与实际需求匹配。'
    )
    
    doc.add_paragraph('算法逻辑：')
    doc.add_paragraph(
        'a) 对每个SKU，汇总近30天内各RDC的订单量\n'
        'b) 计算各RDC需求占比 = 该RDC订单量 / 该SKU总订单量\n'
        'c) 若某RDC需求占比 > 40% 且 该RDC月缺货量 > 50箱\n'
        '   → 建议上调该RDC的分仓分配比例\n'
        'd) 若某RDC需求占比 < 10% 且 该RDC月缺货量 = 0\n'
        '   → 建议下调该RDC的分仓分配比例（释放库存给高需求RDC）'
    )
    
    doc.add_paragraph(
        '输出形式：表格列出建议调整的SKU × RDC组合，包含当前订单占比和月缺货量，'
        '帮助补货计划经理判断是否需要调整分仓策略。'
    )

    add_styled_heading(doc, '3.2 安全库存调整', level=2)
    
    doc.add_paragraph(
        '目标：基于需求波动分析，计算各SKU在各RDC的理论安全库存量，'
        '与当前实际库存对比，给出调整建议。'
    )
    
    doc.add_paragraph('算法逻辑：')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SS_theoretical = 1.65 × σ × √(LT)')
    run.font.size = Pt(12)
    run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(
        '其中：\n'
        '• SS_theoretical = 理论安全库存量（95%服务水平，z=1.65）\n'
        '• σ = 近30天日订单量的标准差\n'
        '• LT = 补货周期 = 4天\n\n'
        '调整建议逻辑：\n'
        '• 若 当前大仓库存 < SS_theoretical × 0.8 → 建议上调安全库存（库存不足）\n'
        '• 若 当前大仓库存 > SS_theoretical × 1.5 → 建议下调安全库存（库存过剩，可释放资金）\n'
        '• 若 CV（变异系数 = σ/μ）> 0.3 → 标记为"高波动SKU"，建议持续重点监控'
    )

    add_styled_heading(doc, '3.3 异常模式检测', level=2)
    
    add_table(doc,
        ['异常模式', '检测条件', '触发阈值', '建议措施'],
        [
            ['需求趋势上涨', '近2周日均缺货量 vs 前2周日均缺货量', '增长 > 50%', '检查是否促销或季节性需求，提前增加补货量'],
            ['满足率持续走低', '同一RDC×SKU连续多天满足率 < 60%', '连续 ≥ 5天', '优先排查库存水位，紧急补充'],
            ['多仓同时缺货', '同一SKU缺货覆盖的RDC数量', '≥ 4个RDC', '总仓库存可能不足，建议紧急调拨或调整生产计划'],
        ]
    )

    doc.add_page_break()

    # ==== Chapter 4: 数据依赖 ====
    add_styled_heading(doc, '4. 数据依赖关系', level=1)
    
    doc.add_paragraph('补货建议模块依赖以下数据表：')
    
    add_table(doc,
        ['数据表', '来源', '关键字段', '用途'],
        [
            ['缺货汇总', 'Sheet1', 'dateStr, materialCode, rdcShortBoxes, dcStock, rdcTransitTotal, abcClass, brand', '当日缺货筛选、库存/在途数据、ABC分类'],
            ['订单明细', 'Sheet5', 'dateStr, warehouse, skuCode, orderQty, totalFulfillQty, firstDayQty, channel', '日均需求计算、满足率计算、RDC需求占比'],
            ['缺货原因维护', 'IndexedDB持久化', '_shortageReasons[date|sku|rdc]', '缺货原因标签匹配'],
            ['品牌月度满足率', 'Sheet2-Block3', 'brand, huanan/huabei/...', '品牌级参考'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        '注：缺货数据对比中维护的原因标签（库存在KA库、低于2/3效期、指定效期发货、入库时间差）'
        '通过 IndexedDB 持久化存储，页面刷新后仍然保留。'
    )

    # ==== Chapter 5: 页面交互 ====
    add_styled_heading(doc, '5. 页面交互设计', level=1)
    
    doc.add_paragraph('页面布局分为上下两大区域：')

    add_table(doc,
        ['区域', '说明'],
        [
            ['🔴 每日补货建议', '位于页面上半部分，显示Top 20重点关注SKU及补货建议。'],
            ['🔵 中期策略调整', '位于页面下半部分，显示分仓需求调整、安全库存调整、异常模式检测结果。'],
            ['RDC筛选', '全局RDC筛选器，可选择只看某个RDC的建议。'],
            ['状态标记', '每条建议可标记为"待处理/已处理"，状态持久化到IndexedDB。'],
            ['导出功能', '一键导出当前筛选条件下的建议为CSV文件。'],
        ]
    )

    # ==== Chapter 6: Appendix ====
    add_styled_heading(doc, '6. 附录：算法公式汇总', level=1)
    
    formulas = [
        ('综合评分 Score', 
         'Score = (1 − 满足率)×0.25 + 缺货趋势×0.15 + ABC系数×0.15 + RDC数×0.10 + 缺货量×0.35'),
        ('日均需求', 'μ = Σ(30天订单量) / 30'),
        ('建议补货量', 'Q = max(0, μ × 4 − S − T)，其中S=大仓库存，T=在途量'),
        ('需求变异系数', 'CV = σ / μ'),
        ('理论安全库存', 'SS = 1.65 × σ × √4'),
        ('调整阈值', '库存 < SS×0.8 → 上调；库存 > SS×1.5 → 下调'),
        ('需求趋势检测', '近2周日均缺货 / 前2周日均缺货 > 1.5 → 趋势上涨'),
        ('满足率预警', '同一RDC×SKU连续5天满足率 < 60% → 异常'),
        ('多仓缺货预警', '同一SKU ≥ 4个RDC同时缺货 → 总仓风险'),
    ]

    for title, formula in formulas:
        p = doc.add_paragraph()
        run = p.add_run(title + '：')
        run.bold = True
        p.add_run(formula)

    # Save
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '补货建议模块设计文档.docx')
    doc.save(output_path)
    print(f'文档已生成：{output_path}')

if __name__ == '__main__':
    main()
